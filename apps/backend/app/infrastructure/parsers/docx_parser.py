"""
===============================================================================
ARCHIVO: docx_parser.py
===============================================================================

CRC CARD (Class)
-------------------------------------------------------------------------------
Clase:
    DocxParser

Responsabilidades:
    - Extraer texto desde DOCX usando python-docx.
    - Incluir contenido de párrafos y tablas (evitar pérdida de datos).
    - Aplicar normalización/truncado (sostenibilidad).
    - Reportar warnings no fatales (p.ej. tablas complejas).

Colaboradores:
    - contracts.ParserOptions / ExtractedText
    - errors.DocumentParsingError / EmptyDocumentError
    - normalize.normalize_text / truncate_text
===============================================================================
"""

from __future__ import annotations

from io import BytesIO

from .contracts import BaseParser, ExtractedText, ParserOptions
from .errors import DocumentParsingError, EmptyDocumentError
from .normalize import normalize_text, truncate_text


class DocxParser(BaseParser):
    """Estrategia de parsing para DOCX (python-docx)."""

    def parse(self, content: bytes, *, options: ParserOptions) -> ExtractedText:
        try:
            from docx import Document
        except ImportError as e:
            raise DocumentParsingError(
                "La librería 'python-docx' no está instalada", original_error=e
            ) from e

        try:
            doc = Document(BytesIO(content))
        except Exception as e:
            raise DocumentParsingError(
                "No se pudo abrir el DOCX (archivo corrupto o inválido)",
                original_error=e,
            ) from e

        parts: list[str] = []
        warnings: list[str] = []

        # 1) Párrafos (caso común)
        try:
            for p in doc.paragraphs:
                t = (p.text or "").strip()
                if t:
                    parts.append(t)
        except Exception as e:
            raise DocumentParsingError(
                "Fallo al leer párrafos del DOCX", original_error=e
            ) from e

        # 2) Tablas (muchos DOCX “serios” guardan contenido aquí)
        try:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            t = (p.text or "").strip()
                            if t:
                                parts.append(t)
        except Exception as e:
            # Senior approach: degradación suave (warning) si falla tablas
            warnings.append(
                f"No se pudo extraer completamente el contenido de tablas: {type(e).__name__}"
            )

        raw_text = "\n".join(parts)

        normalized = normalize_text(
            raw_text, collapse_whitespace=options.normalize_whitespace
        )
        normalized, truncated = truncate_text(normalized, max_chars=options.max_chars)

        result = ExtractedText(
            content=normalized,
            metadata={"source": "docx"},
            warnings=warnings,
            was_truncated=truncated,
        )

        if result.is_empty and not options.allow_empty:
            raise EmptyDocumentError("DOCX sin texto extraíble")

        return result
